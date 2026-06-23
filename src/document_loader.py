"""Document loading and text extraction utilities."""

from __future__ import annotations

from dataclasses import asdict, dataclass
from pathlib import Path
import hashlib
import re
from typing import Iterable


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


class DocumentExtractionError(RuntimeError):
    """Raised when a supported document cannot be extracted."""


@dataclass(frozen=True)
class ExtractedTextRow:
    """Normalized extracted text record used by notebooks and CSV outputs."""

    document_id: str
    source_path: str
    page_number: int | None
    text: str
    word_count: int
    char_count: int
    extraction_method: str
    error: str = ""


def get_extracted_text_columns() -> list[str]:
    """Return the canonical column order for extracted text CSV files."""

    return [
        "document_id",
        "source_path",
        "page_number",
        "text",
        "word_count",
        "char_count",
        "extraction_method",
        "error",
    ]


def count_words(text: str) -> int:
    """Count word-like tokens in extracted text."""

    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text or ""))


def make_document_id(file_path: str | Path, base_dir: str | Path | None = None) -> str:
    """Create a stable document id from a file path."""

    path = Path(file_path)
    id_source = path.resolve()
    if base_dir is not None:
        try:
            id_source = path.resolve().relative_to(Path(base_dir).resolve())
        except ValueError:
            id_source = path.resolve()

    slug = re.sub(r"[^A-Za-z0-9]+", "_", path.stem).strip("_").lower() or "document"
    digest = hashlib.sha1(str(id_source).encode("utf-8")).hexdigest()[:10]
    return f"{slug}_{digest}"


def iter_supported_files(input_dir: str | Path) -> list[Path]:
    """Return supported files below a directory in deterministic order."""

    directory = Path(input_dir)
    if not directory.exists():
        return []
    if not directory.is_dir():
        raise NotADirectoryError(f"Input path is not a directory: {directory}")

    return sorted(
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


def extract_text_from_pdf(
    file_path: str | Path,
    *,
    method: str = "auto",
    base_dir: str | Path | None = None,
) -> list[dict[str, object]]:
    """Extract page-level text from a PDF.

    The default `auto` mode tries PyMuPDF first, then falls back to pdfplumber
    when PyMuPDF fails or returns no non-empty page text.
    """

    path = _validate_file(file_path, ".pdf")
    method = method.lower()
    if method not in {"auto", "pymupdf", "pdfplumber"}:
        raise ValueError("method must be one of: auto, pymupdf, pdfplumber")

    errors: list[str] = []
    first_rows: list[tuple[int, str]] = []

    if method in {"auto", "pymupdf"}:
        try:
            first_rows = _extract_pdf_with_pymupdf(path)
            if method == "pymupdf" or _has_text(first_rows):
                return _build_rows(path, first_rows, "pymupdf", base_dir)
        except Exception as exc:  # pragma: no cover - depends on local PDFs
            if method == "pymupdf":
                raise DocumentExtractionError(f"PyMuPDF failed for {path}: {exc}") from exc
            errors.append(f"PyMuPDF failed: {exc}")

    if method in {"auto", "pdfplumber"}:
        try:
            rows = _extract_pdf_with_pdfplumber(path)
            if method == "pdfplumber" or _has_text(rows) or not first_rows:
                return _build_rows(path, rows, "pdfplumber", base_dir)
        except Exception as exc:  # pragma: no cover - depends on local PDFs
            if method == "pdfplumber":
                raise DocumentExtractionError(f"pdfplumber failed for {path}: {exc}") from exc
            errors.append(f"pdfplumber failed: {exc}")

    if first_rows:
        return _build_rows(path, first_rows, "pymupdf", base_dir)

    detail = "; ".join(errors) if errors else "no pages extracted"
    raise DocumentExtractionError(f"Could not extract PDF text from {path}: {detail}")


def extract_text_from_txt(
    file_path: str | Path,
    *,
    encoding: str = "utf-8",
    base_dir: str | Path | None = None,
) -> list[dict[str, object]]:
    """Extract text from a plain text file."""

    path = _validate_file(file_path, ".txt")
    try:
        text = path.read_text(encoding=encoding)
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")

    return _build_rows(path, [(None, text)], "txt", base_dir)


def extract_text_from_docx(
    file_path: str | Path,
    *,
    base_dir: str | Path | None = None,
) -> list[dict[str, object]]:
    """Extract paragraph and table text from a DOCX file."""

    path = _validate_file(file_path, ".docx")
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise DocumentExtractionError(
            "DOCX extraction requires python-docx. Install it with `pip install python-docx`."
        ) from exc

    document = Document(str(path))
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return _build_rows(path, [(None, "\n".join(chunks))], "docx", base_dir)


def extract_text_from_file(
    file_path: str | Path,
    *,
    base_dir: str | Path | None = None,
) -> list[dict[str, object]]:
    """Extract text from one supported file."""

    path = Path(file_path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path, base_dir=base_dir)
    if suffix == ".txt":
        return extract_text_from_txt(path, base_dir=base_dir)
    if suffix == ".docx":
        return extract_text_from_docx(path, base_dir=base_dir)

    raise ValueError(f"Unsupported file type: {path.suffix}")


def extract_documents_from_directory(input_dir: str | Path) -> list[dict[str, object]]:
    """Extract text from all supported files in a directory.

    Extraction errors are returned as rows so one bad file does not stop a full
    notebook run.
    """

    directory = Path(input_dir)
    records: list[dict[str, object]] = []

    for path in iter_supported_files(directory):
        try:
            records.extend(extract_text_from_file(path, base_dir=directory))
        except Exception as exc:  # pragma: no cover - defensive batch behavior
            records.append(_error_row(path, directory, str(exc)))

    return records


def _validate_file(file_path: str | Path, expected_suffix: str) -> Path:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    if not path.is_file():
        raise ValueError(f"Path is not a file: {path}")
    if path.suffix.lower() != expected_suffix:
        raise ValueError(f"Expected {expected_suffix} file, got: {path.suffix}")
    return path


def _extract_pdf_with_pymupdf(path: Path) -> list[tuple[int, str]]:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise DocumentExtractionError(
            "PDF extraction with PyMuPDF requires PyMuPDF. Install it with `pip install PyMuPDF`."
        ) from exc

    rows: list[tuple[int, str]] = []
    document = fitz.open(path)
    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            rows.append((page_index + 1, page.get_text("text") or ""))
    finally:
        document.close()
    return rows


def _extract_pdf_with_pdfplumber(path: Path) -> list[tuple[int, str]]:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - environment dependent
        raise DocumentExtractionError(
            "PDF extraction with pdfplumber requires pdfplumber. Install it with `pip install pdfplumber`."
        ) from exc

    rows: list[tuple[int, str]] = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            rows.append((page_index, page.extract_text() or ""))
    return rows


def _build_rows(
    path: Path,
    page_texts: Iterable[tuple[int | None, str]],
    extraction_method: str,
    base_dir: str | Path | None,
) -> list[dict[str, object]]:
    document_id = make_document_id(path, base_dir)
    source_path = str(path)
    rows = [
        asdict(
            ExtractedTextRow(
                document_id=document_id,
                source_path=source_path,
                page_number=page_number,
                text=text or "",
                word_count=count_words(text or ""),
                char_count=len(text or ""),
                extraction_method=extraction_method,
            )
        )
        for page_number, text in page_texts
    ]

    if not rows:
        rows.append(
            asdict(
                ExtractedTextRow(
                    document_id=document_id,
                    source_path=source_path,
                    page_number=None,
                    text="",
                    word_count=0,
                    char_count=0,
                    extraction_method=extraction_method,
                    error="No text records extracted.",
                )
            )
        )
    return rows


def _error_row(path: Path, base_dir: str | Path | None, error: str) -> dict[str, object]:
    return asdict(
        ExtractedTextRow(
            document_id=make_document_id(path, base_dir),
            source_path=str(path),
            page_number=None,
            text="",
            word_count=0,
            char_count=0,
            extraction_method="error",
            error=error,
        )
    )


def _has_text(rows: Iterable[tuple[int | None, str]]) -> bool:
    return any((text or "").strip() for _, text in rows)
